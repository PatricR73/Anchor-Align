"""S6 — transcript ingest: DOCX/TXT parsing, page-number and header/footer
cleanup. PDF is out of scope.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

import docx

from anchor_align.exceptions import IngestError
from anchor_align.models import EditedToken

logger = logging.getLogger(__name__)

_SENTENCE_END = re.compile(r"[.!?]$")
_TOKEN_RE = re.compile(r"\S+")
_PAGE_NUMBER_ONLY = re.compile(r"^\d+$")


def _extract_text(path: Path) -> str:
    """Header/footer text lives outside python-docx's `document.paragraphs`
    (a separate XML part), so it's excluded by construction here, not by
    filtering. A body paragraph that's purely digits — a page number that
    leaked into the body instead of a real header/footer — is dropped as a
    cheap heuristic; anything with real words survives.
    """
    if path.suffix.lower() == ".docx":
        document = docx.Document(str(path))
        paragraphs = [
            p.text
            for p in document.paragraphs
            if p.text.strip() and not _PAGE_NUMBER_ONLY.match(p.text.strip())
        ]
        return "\n".join(paragraphs)
    return path.read_text(encoding="utf-8")


def parse_transcript(path: Path) -> list[EditedToken]:
    """Return cleaned tokens from a .docx or .txt transcript. Implements
    TranscriptExtractor.extract.

    `char_offset` is the offset into the extracted text this function
    builds internally (post header/footer/page-number removal, paragraphs
    joined with "\\n") — not an offset into the original .docx's raw XML,
    which has no single meaningful linear character axis to offset into.
    """
    try:
        text = _extract_text(path)
    except (OSError, ValueError) as e:
        raise IngestError(f"could not read transcript {path}: {e}") from e

    tokens: list[EditedToken] = []
    sentence_id = 0
    for index, m in enumerate(_TOKEN_RE.finditer(text)):
        word = m.group()
        is_end = bool(_SENTENCE_END.search(word))
        tokens.append(
            EditedToken(
                text=word,
                index=index,
                char_offset=m.start(),
                sentence_id=sentence_id,
                is_sentence_end=is_end,
            )
        )
        if is_end:
            sentence_id += 1
    logger.info("parsed %d tokens from %s", len(tokens), path)
    return tokens
