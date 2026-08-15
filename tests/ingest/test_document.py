"""S6 — transcript ingest: DOCX/TXT parsing, page-number/header/footer
cleanup."""

from __future__ import annotations

import docx

from anchor_align.ingest.document import parse_transcript


def test_parse_txt_produces_tokens_with_correct_offsets(tmp_path):
    path = tmp_path / "transcript.txt"
    path.write_text("Hello world. Second sentence!", encoding="utf-8")

    tokens = parse_transcript(path)

    assert [t.text for t in tokens] == ["Hello", "world.", "Second", "sentence!"]
    assert tokens[0].char_offset == 0
    assert tokens[1].char_offset == 6
    assert tokens[0].index == 0
    assert tokens[3].index == 3


def test_sentence_ids_increment_on_sentence_end(tmp_path):
    path = tmp_path / "t.txt"
    path.write_text("One two. Three four? Five.", encoding="utf-8")

    tokens = parse_transcript(path)

    assert [t.sentence_id for t in tokens] == [0, 0, 1, 1, 2]
    assert [t.is_sentence_end for t in tokens] == [False, True, False, True, True]


def test_parse_docx_extracts_paragraph_text(tmp_path):
    path = tmp_path / "transcript.docx"
    document = docx.Document()
    document.add_paragraph("Hello world.")
    document.add_paragraph("Second paragraph here.")
    document.save(str(path))

    tokens = parse_transcript(path)

    assert [t.text for t in tokens] == ["Hello", "world.", "Second", "paragraph", "here."]


def test_parse_docx_drops_page_number_only_paragraphs(tmp_path):
    path = tmp_path / "transcript.docx"
    document = docx.Document()
    document.add_paragraph("Real content here.")
    document.add_paragraph("42")  # a page number that leaked into the body
    document.add_paragraph("More real content.")
    document.save(str(path))

    tokens = parse_transcript(path)
    text = [t.text for t in tokens]

    assert "42" not in text
    assert "Real" in text
    assert "More" in text


def test_parse_docx_headers_and_footers_are_not_included(tmp_path):
    path = tmp_path / "transcript.docx"
    document = docx.Document()
    document.sections[0].header.paragraphs[0].text = "Confidential Header"
    document.sections[0].footer.paragraphs[0].text = "Page footer text"
    document.add_paragraph("Only this should appear.")
    document.save(str(path))

    tokens = parse_transcript(path)
    text = " ".join(t.text for t in tokens)

    assert "Confidential" not in text
    assert "footer" not in text
    assert "Only" in text


def test_parse_docx_skips_empty_paragraphs(tmp_path):
    path = tmp_path / "transcript.docx"
    document = docx.Document()
    document.add_paragraph("First.")
    document.add_paragraph("")  # blank paragraph, common in real docs
    document.add_paragraph("Second.")
    document.save(str(path))

    tokens = parse_transcript(path)

    assert [t.text for t in tokens] == ["First.", "Second."]


def test_empty_txt_yields_no_tokens(tmp_path):
    path = tmp_path / "empty.txt"
    path.write_text("", encoding="utf-8")
    assert parse_transcript(path) == []
